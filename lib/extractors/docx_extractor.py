"""DOCX extractor using python-docx."""

from typing import Any
from docx import Document


def extract_from_docx(filepath: str) -> dict[str, Any]:
    """Extract structured content from DOCX."""
    doc = Document(filepath)
    sections = []

    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            level = _extract_heading_level(para.style.name)
            sections.append({
                "level": level,
                "title": para.text,
                "type": "heading",
            })
        elif para.text.strip():
            # Add content to last section or create initial section
            if sections and sections[-1].get("type") != "heading":
                sections[-1]["content"] = sections[-1].get("content", "") + "\n" + para.text
            elif sections:
                sections.append({
                    "content": para.text,
                    "parent_heading": sections[-1]["title"] if sections else None,
                })
            else:
                sections.append({
                    "content": para.text,
                    "parent_heading": None,
                })

    # Also extract tables
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)

    return {
        "file": filepath,
        "sections": sections,
        "tables": tables,
    }


def _extract_heading_level(style_name: str) -> int:
    """Extract heading level from style name."""
    if "Heading" in style_name:
        try:
            return int(style_name.replace("Heading ", ""))
        except ValueError:
            return 1
    elif "Title" in style_name:
        return 0
    return 1